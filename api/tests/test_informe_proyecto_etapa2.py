from io import BytesIO
from unittest.mock import MagicMock, patch

from django.test import SimpleTestCase
from docx import Document

from api.informe_word_service import (
    _add_matriz_bruta_scientific_explanation,
    _begin_table_block,
    _compute_column_widths_cm,
    _display_valor_evaluacion,
    _estimate_evaluation_table_height_cm,
    _estimate_table_height_cm,
    _evaluacion_matrix_headers,
    _evaluacion_matrix_rows,
    _page_used_after_block,
    build_informe_proyecto_docx,
)


class InformeProyectoEtapa2Tests(SimpleTestCase):
    def test_display_valor_riesgo(self):
        col = {'input_kind': 'riesgo'}
        self.assertEqual(_display_valor_evaluacion('0.3|0.7', col), 'P=0.3 · C=0.7')

    def test_evaluacion_matrix_alternativas_en_x_y_terminales_en_y(self):
        alt_a = MagicMock()
        alt_a.id = 1
        alt_a.nombre = 'Alt A'
        alt_b = MagicMock()
        alt_b.id = 2
        alt_b.nombre = 'Alt B'
        columnas = [
            {'key': 'k1', 'terminal_nombre': 'Rango'},
            {'key': 'k2', 'terminal_nombre': 'Autonomía'},
        ]
        alternativas_valores = [
            (alt_a, {'k1': '100', 'k2': ''}),
            (alt_b, {'k1': '80', 'k2': '12'}),
        ]

        headers = _evaluacion_matrix_headers(alternativas_valores)
        rows = _evaluacion_matrix_rows(alternativas_valores, columnas)

        self.assertEqual(headers, ['Nodo terminal', 'Alt A', 'Alt B'])
        self.assertEqual(rows[0], ['Rango', '100', '80'])
        self.assertEqual(rows[1], ['Autonomía', '—', '12'])

    def test_dos_tablas_cortas_caben_en_una_pagina(self):
        headers = [
            'Nodo terminal',
            'Ship_A',
            'Ship_B',
            'Ship_C',
            'Ship_D',
            'Ship_E',
            'Ship_F',
        ]
        rows = [
            ['Acquisition Cost', '140.88', '140.88', '140.88', '140.88', '140.88', '140.88'],
            ['Operation Cost', '17.01', '17.01', '17.01', '17.01', '17.01', '17.01'],
            ['Maintenance Cost', '2.48', '2.48', '2.48', '2.48', '2.48', '2.48'],
        ]
        first = _estimate_evaluation_table_height_cm(
            headers,
            rows,
            extra_heading_count=3,
        )
        second = _estimate_evaluation_table_height_cm(
            headers,
            rows,
            extra_heading_count=1,
        )
        self.assertLessEqual(2.8 + first + second + 0.9, 17.5)

    def test_residuo_tras_tabla_casi_pagina_llena_no_es_fantasma(self):
        # 18.3 cm ≈ 1 página + 0.8: en Word suele caber en una hoja.
        # No inventar used=0.8 (eso hacía que Cost se partiera bajo Effectiveness).
        self.assertEqual(_page_used_after_block(0.0, 18.3), 17.5)
        # Tabla claramente multipágina (≥2 páginas + residuo útil).
        self.assertAlmostEqual(
            _page_used_after_block(0.0, 17.5 * 2 + 3.0),
            3.0,
            places=5,
        )

    def test_anchos_columna_priorizan_contenido_largo(self):
        headers = [
            'Escenario',
            'Dimensión',
            'Grupo de afinidad',
            'Nodo intermedio — MOB',
            'Nodo terminal — DT',
            'Descripción o función',
        ]
        rows = [
            [
                'Escenario base',
                'Cost',
                '—',
                '—',
                'Acquisition Cost',
                'Función: razon_inversa · L=250, U=50',
            ],
            [
                'Escenario base',
                'Cost',
                '—',
                '—',
                'Operation Cost',
                'Función: razon_inversa · L=40, U=5',
            ],
            [
                'Escenario base',
                'Cost',
                '—',
                '—',
                'Maintenance Cost',
                'Función: razon_inversa · L=10, U=1',
            ],
        ]
        widths = _compute_column_widths_cm(headers, rows)
        self.assertEqual(len(widths), 6)
        self.assertAlmostEqual(sum(widths), 16.8, places=2)
        # Columnas con "—" deben ser más estrechas que descripción/terminal.
        self.assertLess(widths[2], widths[4])
        self.assertLess(widths[3], widths[5])
        self.assertGreater(widths[5], widths[1])

        # Dos tablas cortas Cost/Risk deben poder compartir página.
        height = _estimate_table_height_cm(headers, rows, extra_heading_count=1)
        self.assertLess(height, 7.0)
        self.assertLessEqual(2.8 + height + height + 0.9, 17.5)

    def test_begin_table_block_usa_residuo_si_la_siguiente_cabe_completa(self):
        doc = Document()
        # Simula una tabla larga que dejó 0.8 cm ocupados en la página actual.
        setattr(doc, '_report_table_page_used_cm', 0.8)
        _begin_table_block(doc, estimated_height_cm=5.0)
        # No debe forzar salto: 0.8 + 5.0 + buffer cabe y el residuo sigue útil.
        self.assertAlmostEqual(
            getattr(doc, '_report_table_page_used_cm'),
            5.8,
            places=5,
        )
        # Si no cabe completa, sí salta; tras colocar 13 cm el residuo (<7.5) llena la página.
        _begin_table_block(doc, estimated_height_cm=13.0)
        self.assertEqual(getattr(doc, '_report_table_page_used_cm'), 17.5)

    def test_begin_table_block_no_parte_tabla_en_espacio_insuficiente(self):
        """Si no cabe completa (título+tabla), nueva página — sin excepciones."""
        doc = Document()
        # Página casi llena tras una tabla previa.
        setattr(doc, '_report_table_page_used_cm', 12.0)
        body_before = len(doc.element.body)
        _begin_table_block(doc, estimated_height_cm=8.0)
        # Debe insertar salto: 12 + 8 + buffer > 17.5.
        self.assertGreater(len(doc.element.body), body_before)
        self.assertAlmostEqual(
            getattr(doc, '_report_table_page_used_cm'),
            8.0,
            places=5,
        )

    def test_dos_tablas_jerarquicas_cortas_comparten_pagina(self):
        doc = Document()
        setattr(doc, '_report_table_page_used_cm', 2.8)
        headers = [
            'Escenario', 'Dimensión', 'Grupo de afinidad',
            'Nodo intermedio — MOB', 'Nodo terminal — DT', 'Descripción o función',
        ]
        rows = [
            ['Escenario base', 'Cost', '—', '—', 'Acquisition Cost', 'Función: razon_inversa · L=250, U=50'],
            ['Escenario base', 'Cost', '—', '—', 'Operation Cost', 'Función: razon_inversa · L=40, U=5'],
            ['Escenario base', 'Cost', '—', '—', 'Maintenance Cost', 'Función: razon_inversa · L=10, U=1'],
        ]
        h = _estimate_table_height_cm(headers, rows, extra_heading_count=1)
        body_before = len(doc.element.body)
        _begin_table_block(doc, h)
        _begin_table_block(doc, h)
        # Segunda tabla corta debe quedar en la misma página (sin salto).
        self.assertEqual(len(doc.element.body), body_before)
        # Tras dos tablas el residuo es pequeño → página marcada llena.
        self.assertEqual(getattr(doc, '_report_table_page_used_cm'), 17.5)

    def test_tras_tabla_grande_no_arranca_otra_en_residuo_chico(self):
        doc = Document()
        setattr(doc, '_report_table_page_used_cm', 0.0)
        # Tabla grande (~14 cm): deja poco residuo.
        _begin_table_block(doc, estimated_height_cm=14.0)
        self.assertEqual(getattr(doc, '_report_table_page_used_cm'), 17.5)
        body_before = len(doc.element.body)
        # Cost corta no debe empezar en ese residuo.
        _begin_table_block(doc, estimated_height_cm=5.8)
        self.assertGreater(len(doc.element.body), body_before)
        self.assertAlmostEqual(getattr(doc, '_report_table_page_used_cm'), 5.8, places=5)

    def test_tabla_casi_pagina_no_deja_residuo_fantasma_para_cost(self):
        """Regresión captura: Effectiveness ~1 pág estimada 17.8 → Cost partida.

        El overflow módulo página inventaba ~0.3 cm usados en "página 2";
        Cost creía caber ahí y Word la partía bajo Effectiveness.
        """
        from api.informe_word_service import _page_used_after_block

        # Estimación apenas > página usable (caso real Effectiveness M6).
        self.assertEqual(_page_used_after_block(0.0, 17.85), 17.5)
        self.assertEqual(_page_used_after_block(0.0, 19.0), 17.5)

        doc = Document()
        setattr(doc, '_report_table_page_used_cm', 0.0)
        _begin_table_block(doc, estimated_height_cm=17.85)
        self.assertEqual(getattr(doc, '_report_table_page_used_cm'), 17.5)
        body_before = len(doc.element.body)
        _begin_table_block(doc, estimated_height_cm=5.8)
        self.assertGreater(len(doc.element.body), body_before)
        self.assertAlmostEqual(getattr(doc, '_report_table_page_used_cm'), 5.8, places=5)

    def test_explicacion_cientifica_estilo_articulo(self):
        leaf_range = {
            'kind': 'leaf',
            'nombre': 'Range',
            'valor': 0.443324,
            'familia_label': 'Min-max',
            'utilidad_tipo': 'LinearUtilityFunction',
            'parametros': {'L': 1000, 'U': 6000},
            'debug': {'L': 1000, 'U': 6000, 'is_increasing': True},
            'escenarios': [
                {'nombre': 'M1', 'x': '4197.133992', 'u': 0.639427, 'peso': 20},
                {'nombre': 'M2', 'x': '1468.476199', 'u': 0.093695, 'peso': 20},
            ],
            'hijos': [],
        }
        leaf_maneuver = {
            'kind': 'leaf',
            'nombre': 'Maneuverability',
            'valor': 0.5,
            'familia_label': 'Discreta',
            'utilidad_tipo': 'DiscreteUtilityFunction',
            'parametros': {},
            'debug': {},
            'escenarios': [
                {'nombre': 'M1', 'x': 'Medium', 'u': 0.5, 'peso': 20},
                {'nombre': 'M2', 'x': 'High', 'u': 1.0, 'peso': 20},
            ],
            'hijos': [],
        }
        mobility = {
            'kind': 'rollup',
            'nombre': 'Mobility',
            'valor': 0.466643,
            'suma_pesos': 100,
            'hijos': [
                {'nombre': 'Range', 'valor': 0.443324, 'peso': 50, 'trace': leaf_range},
                {
                    'nombre': 'Maneuverability',
                    'valor': 0.5,
                    'peso': 50,
                    'trace': leaf_maneuver,
                },
            ],
        }
        trace = {
            'kind': 'dimension',
            'nombre': 'Effectiveness',
            'valor': 0.474107,
            'modo_valor_terminal': 'utilidad',
            'suma_pesos': 100,
            'hijos': [
                {
                    'nombre': 'Mobility',
                    'valor': 0.466643,
                    'peso': 100,
                    'trace': mobility,
                },
            ],
        }
        resultados = [{
            'id': 1,
            'nombre': 'Ship_A',
            'dimensiones': [{
                'omoe_id': 1,
                'omoe_nombre': 'Effectiveness',
                'escenario_agregacion': 'compensatorio',
                'valor': 0.474107,
                'detalle': {'metodo': 'nodo_arbol', 'trace': trace},
            }],
        }]
        document = Document()

        _add_matriz_bruta_scientific_explanation(document, resultados)

        text = '\n'.join(p.text for p in document.paragraphs)
        table_text = '\n'.join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn(
            'Ejemplo numérico: cálculo de Effectiveness para la alternativa Ship_A',
            text,
        )
        self.assertIn('Paso 1. Transformación de los nodos terminales', text)
        self.assertIn('Paso 2. Agregación jerárquica', text)
        self.assertIn('Paso 3. Asignación del resultado en la matriz general', text)
        self.assertIn('Multi-Attribute Value Theory', text)
        # Ecuaciones van como OMML nativo (no texto plano ni imagen).
        omml_xml = '\n'.join(p._p.xml for p in document.paragraphs)
        self.assertIn('oMath', omml_xml)
        self.assertIn('Range', omml_xml)
        self.assertIn('Mobility', omml_xml)
        self.assertIn('Ship_A', omml_xml)
        self.assertIn('0.474107', omml_xml)
        self.assertNotIn('w:drawing', omml_xml)
        self.assertIn('Medium', table_text)
        self.assertIn('4197.133992', table_text)
        self.assertLessEqual(len(document.tables), 3)

    @patch('api.informe_word_service.build_evaluacion_schema')
    @patch('api.informe_word_service._load_alternativas_valores')
    @patch('api.informe_word_service._add_etapa1_arboles_section')
    @patch('api.informe_word_service._add_alternativas_section')
    @patch('api.informe_word_service._add_etapa1_estructura_jerarquica')
    @patch('api.informe_word_service._add_etapa1_pesos_nodo')
    @patch('api.informe_word_service._build_matriz_bruta_calculo')
    def test_docx_incluye_etapa2_y_matriz_bruta_calculo(
        self,
        mock_build_bruta,
        _mock_pesos,
        _mock_estructura,
        _mock_alternativas,
        _mock_arboles,
        mock_load_vals,
        mock_schema,
    ):
        mock_build_bruta.return_value = (
            ['Alternativa', 'Cost', 'Effectiveness'],
            [
                ['Sentido', 'Minimizar', 'Maximizar'],
                ['Patrullera A', 0.42, 0.81],
            ],
            None,
        )
        mock_schema.return_value = {
            'dimensiones': [
                {
                    'omoe_id': 10,
                    'omoe_nombre': 'Cost',
                    'escenarios': [{'id': 1, 'nombre': 'Escenario base'}],
                    'columnas': [
                        {
                            'key': 'nodo:1:1',
                            'omoe_id': 10,
                            'escenario_id': 1,
                            'escenario_nombre': 'Escenario base',
                            'terminal_nombre': 'Acquisition Cost',
                        },
                    ],
                },
            ],
            'columnas': [],
        }
        alt = MagicMock()
        alt.id = 1
        alt.nombre = 'Patrullera A'
        mock_load_vals.return_value = [(alt, {'nodo:1:1': '1200'})]

        proyecto = MagicMock()
        proyecto.nombre = 'Demo'
        proyecto.descripcion = ''
        proyecto.eslora_maxima = ''
        proyecto.desplazamiento = ''
        proyecto.velocidad_maxima = ''
        proyecto.velocidad_crucero = ''
        proyecto.tripulacion = ''
        proyecto.autonomia = ''
        proyecto.propulsion = ''
        proyecto.posicionamiento_dinamico = ''
        proyecto.laboratorios = ''
        proyecto.otras_caracteristicas = ''

        progress_updates = []
        content = build_informe_proyecto_docx(
            proyecto,
            progress_callback=lambda percent, stage: progress_updates.append(
                (percent, stage)
            ),
        )
        document = Document(BytesIO(content))
        text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        table_text = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        self.assertIn('Etapa 2. Matrices de evaluación', text)
        self.assertIn('2.3. Metodología del cálculo de la matriz general', text)
        self.assertIn('2.4. Matriz general o matriz inicial', text)
        self.assertIn('Formulación general', text)
        self.assertIn('Escenario: Escenario base', text)
        self.assertIn('Patrullera A', table_text)
        self.assertIn('1200', table_text)
        self.assertIn('Matriz bruta del cálculo', text)
        self.assertIn('Cost', table_text)
        self.assertIn('0.4200', table_text)
        self.assertTrue(
            any('tablas de evaluación' in stage for _, stage in progress_updates)
        )
