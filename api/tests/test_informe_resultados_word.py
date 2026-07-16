"""Tests del Informe de resultados (por cálculo)."""
from __future__ import annotations

from io import BytesIO

from django.test import SimpleTestCase
from docx import Document

from api.informe_resultados_word_service import (
    _dimensiones_from_resultado,
    _pareto_inputs,
    build_informe_resultados_docx,
)
from api.models import Proyecto
from api.pareto_solver import ParetoSolver


class InformeResultadosWordTests(SimpleTestCase):
    def test_build_informe_resultados_docx_basico(self):
        proyecto = Proyecto(id=1, nombre='Proyecto prueba')
        resultado = {
            'historial_id': 99,
            'opciones_calculo': {
                'aplicar_pareto': True,
                'pareto_epsilon': 0.0,
                'normalizacion_metodo': 'directional_minmax',
                'metodo_pesos': 'equal_weights',
                'metodo_madm': 'topsis',
                'dimensiones_normalizar': ['Effectiveness', 'Cost', 'Risk'],
            },
            'alternativas': [
                {
                    'id': 1,
                    'nombre': 'Ship A',
                    'apodo': 'A',
                    'ranking': 1,
                    'score_madm': 0.82,
                    'excluida_pareto': False,
                    'dimensiones': [
                        {
                            'omoe_id': 10,
                            'omoe_nombre': 'Effectiveness',
                            'rama_evaluacion': 'omoe',
                            'valor': 0.7,
                        },
                        {
                            'omoe_id': 20,
                            'omoe_nombre': 'Cost',
                            'rama_evaluacion': 'omoc',
                            'valor': 0.4,
                        },
                        {
                            'omoe_id': 30,
                            'omoe_nombre': 'Risk',
                            'rama_evaluacion': 'omor',
                            'valor': 0.3,
                        },
                    ],
                },
                {
                    'id': 2,
                    'nombre': 'Ship B',
                    'apodo': 'B',
                    'ranking': 2,
                    'score_madm': 0.55,
                    'excluida_pareto': False,
                    'dimensiones': [
                        {
                            'omoe_id': 10,
                            'omoe_nombre': 'Effectiveness',
                            'rama_evaluacion': 'omoe',
                            'valor': 0.5,
                        },
                        {
                            'omoe_id': 20,
                            'omoe_nombre': 'Cost',
                            'rama_evaluacion': 'omoc',
                            'valor': 0.6,
                        },
                        {
                            'omoe_id': 30,
                            'omoe_nombre': 'Risk',
                            'rama_evaluacion': 'omor',
                            'valor': 0.5,
                        },
                    ],
                },
            ],
            'pareto': {
                'pareto_alternatives': ['Ship A', 'Ship B'],
                'dominated_alternatives': [],
            },
            'normalizacion': {
                'method': 'directional_minmax',
                'dimensions': ['Effectiveness', 'Cost', 'Risk'],
                'pareto_alternatives': ['Ship A', 'Ship B'],
                'normalized_matrix': [[1.0, 1.0, 1.0], [0.0, 0.0, 0.0]],
            },
            'pesos': {
                'method': 'equal_weights',
                'weights': [1 / 3, 1 / 3, 1 / 3],
                'weights_by_dimension': {
                    'Effectiveness': 1 / 3,
                    'Cost': 1 / 3,
                    'Risk': 1 / 3,
                },
            },
            'madm': {
                'method': 'topsis',
                'best_alternative': 'Ship A',
            },
            'matriz_original': [[0.7, 0.4, 0.3], [0.5, 0.6, 0.5]],
        }

        content = build_informe_resultados_docx(
            proyecto,
            resultado=resultado,
            nombre_calculo='Cálculo demo TOPSIS',
        )
        self.assertGreater(len(content), 2000)

        doc = Document(BytesIO(content))
        text = '\n'.join(p.text for p in doc.paragraphs)
        # Also scan table cells for trace headings that may live only in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += '\n' + cell.text

        self.assertIn('Informe de resultados', text)
        self.assertIn('Filtro de Pareto', text)
        self.assertIn('Normalización', text)
        self.assertIn('TOPSIS', text)
        self.assertIn('Gráficos de resultados', text)
        self.assertTrue(
            'Desarrollo numérico' in text or 'Sustitución' in text,
            'Expected numerical development / substitution trace in document',
        )
        self.assertTrue(
            'Gráficos unidimensionales' in text or '3.7.1' in text,
            'Expected 1D charts section',
        )
        self.assertTrue(
            'tridimensionales' in text or '3.7.3' in text,
            'Expected 3D charts section',
        )


class ParetoInputsMatchInterfaceTests(SimpleTestCase):
    """El Pareto del informe debe operar sobre el mismo subconjunto de dimensiones
    y con las mismas direcciones que el pipeline/interfaz."""

    def _resultado_subset(self):
        # 3 dimensiones en cada alternativa, pero el cálculo solo usó 2 (activas).
        # En las dos dimensiones activas, A domina a B; en la inactiva B es mejor.
        return {
            'opciones_calculo': {
                'aplicar_pareto': True,
                # Solo dos dimensiones activas (la interfaz excluyó "Extra").
                'dimensiones_normalizar': ['Effectiveness', 'Cost'],
                'direcciones_por_dimension': {
                    'Effectiveness': 'max',
                    'Cost': 'min',
                    'Extra': 'max',
                },
            },
            'alternativas': [
                {
                    'id': 1, 'nombre': 'A', 'excluida_pareto': False,
                    'dimensiones': [
                        {'omoe_id': 10, 'omoe_nombre': 'Effectiveness',
                         'rama_evaluacion': 'omoe', 'valor': 0.9},
                        {'omoe_id': 20, 'omoe_nombre': 'Cost',
                         'rama_evaluacion': 'omoc', 'valor': 0.2},
                        {'omoe_id': 30, 'omoe_nombre': 'Extra',
                         'rama_evaluacion': 'omoe', 'valor': 0.1},
                    ],
                },
                {
                    'id': 2, 'nombre': 'B', 'excluida_pareto': True,
                    'dimensiones': [
                        {'omoe_id': 10, 'omoe_nombre': 'Effectiveness',
                         'rama_evaluacion': 'omoe', 'valor': 0.5},
                        {'omoe_id': 20, 'omoe_nombre': 'Cost',
                         'rama_evaluacion': 'omoc', 'valor': 0.6},
                        {'omoe_id': 30, 'omoe_nombre': 'Extra',
                         'rama_evaluacion': 'omoe', 'valor': 0.9},
                    ],
                },
            ],
            'pareto': {
                'pareto_alternatives': ['A'],
                'dominated_alternatives': ['B'],
            },
        }

    def test_pareto_inputs_only_active_dimensions(self):
        resultado = self._resultado_subset()
        opciones = resultado['opciones_calculo']
        dims = _dimensiones_from_resultado(resultado)

        alt_names, matrix, dim_names, directions = _pareto_inputs(
            resultado, dims, opciones,
        )

        self.assertEqual(dim_names, ['Effectiveness', 'Cost'])
        self.assertEqual(directions, ['max', 'min'])
        self.assertEqual(alt_names, ['A', 'B'])
        self.assertEqual(matrix, [[0.9, 0.2], [0.5, 0.6]])

    def test_report_pareto_matches_interface_exclusion(self):
        resultado = self._resultado_subset()
        opciones = resultado['opciones_calculo']
        dims = _dimensiones_from_resultado(resultado)

        alt_names, matrix, dim_names, directions = _pareto_inputs(
            resultado, dims, opciones,
        )
        result = ParetoSolver(
            matrix=matrix,
            dimensions=dim_names,
            directions=directions,
            alternatives=alt_names,
        ).solve()

        # Debe coincidir con lo guardado por la interfaz: B queda dominada.
        self.assertEqual(
            sorted(result.dominated_alternatives),
            sorted(resultado['pareto']['dominated_alternatives']),
        )
        self.assertEqual(
            sorted(result.pareto_alternatives),
            sorted(resultado['pareto']['pareto_alternatives']),
        )
