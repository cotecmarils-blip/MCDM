"""Informe Word de resultados por cálculo (Etapa 3 / simulación).

Mismo formato visual que el Informe de proyecto (logos, tipografía, tablas),
pero contenido paso a paso del pipeline MADM del cálculo seleccionado.
"""
from __future__ import annotations

import itertools
from io import BytesIO
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, RGBColor

from .models import Proyecto, SimulacionHistorial
from .simulacion_service import (
    _MADM_LABELS,
    _PESOS_LABELS,
    resumen_opciones_calculo,
)
from .informe_math import add_eq_line
from .informe_word_service import (
    _add_header_logos,
    _add_heading,
    _add_table,
    _fmt,
    _set_run_font,
    _setup_document_fonts,
)
from .informe_resultados_trace import (
    _fmt_num,
    add_madm_trace,
    add_normalization_trace,
    add_pareto_trace,
    add_picture_block,
    add_weights_trace,
    bar_1d_png,
    extract_decision_matrix,
    radar_png,
    ranking_bar_png,
    ranking_util_madm_png,
    resolve_directions_for_dims,
    scatter_2d_heat_png,
    scatter_2d_png,
    scatter_3d_png,
)

NORMALIZATION_DOCS: dict[str, dict[str, Any]] = {
    'directional_minmax': {
        'title': 'Min-max direccional',
        'intro': (
            'Escala cada columna al intervalo [0, 1] respetando si el criterio '
            'es de beneficio (max) o costo (min).'
        ),
        'equations': [
            ('Beneficio (max)', r'r_{ij}=\dfrac{x_{ij}-\min_j}{\max_j-\min_j}'),
            ('Costo (min)', r'r_{ij}=\dfrac{\max_j-x_{ij}}{\max_j-\min_j}'),
        ],
    },
    'vector': {
        'title': 'Vectorial',
        'intro': 'Divide cada valor por la norma euclídea de su columna.',
        'equations': [
            ('Fórmula', r'r_{ij}=\dfrac{x_{ij}}{\sqrt{\sum_i x_{ij}^{2}}}'),
        ],
    },
    'directional_vector': {
        'title': 'Vectorial direccional',
        'intro': 'Normalización vectorial con orientación; en costos se complementa a 1.',
        'equations': [
            ('Beneficio (max)', r'r_{ij}=\dfrac{x_{ij}}{\sqrt{\sum_i x_{ij}^{2}}}'),
            ('Costo (min)', r'r_{ij}=1-\dfrac{x_{ij}}{\sqrt{\sum_i x_{ij}^{2}}}'),
        ],
    },
    'sum': {
        'title': 'Por suma',
        'intro': 'Proporción respecto a la suma de la columna; en costos usa el inverso.',
        'equations': [
            ('Beneficio (max)', r'r_{ij}=\dfrac{x_{ij}}{\sum_i x_{ij}}'),
            ('Costo (min)', r'r_{ij}=\dfrac{1/x_{ij}}{\sum_i (1/x_{ij})}'),
        ],
    },
}

WEIGHT_DOCS: dict[str, dict[str, Any]] = {
    'equal_weights': {
        'title': 'Pesos iguales',
        'intro': 'Cada dimensión recibe el mismo peso.',
        'equations': [
            ('Fórmula', r'w_{j}=\dfrac{1}{m}'),
        ],
    },
    'user_defined_weights': {
        'title': 'Pesos definidos por el usuario',
        'intro': (
            'El usuario indica porcentajes pⱼ (%) que deben sumar 100; '
            'el sistema los convierte a pesos normalizados.'
        ),
        'equations': [
            ('Fórmula', r'w_{j}=\dfrac{p_{j}}{\sum_k p_{k}}'),
            ('Restricción', r'\sum_j p_{j}=100\,\%'),
        ],
    },
    'entropy': {
        'title': 'Entropía',
        'intro': 'Mayor peso a dimensiones con mayor diversidad de información.',
        'equations': [
            ('Proporción', r'p_{ij}=\dfrac{x_{ij}}{\sum_i x_{ij}}'),
            ('Entropía', r'e_{j}=-\dfrac{\sum_i p_{ij}\,\ln(p_{ij})}{\ln(n)}'),
            ('Peso', r'w_{j}=\dfrac{1-e_{j}}{\sum_k (1-e_{k})}'),
        ],
    },
    'critic': {
        'title': 'CRITIC',
        'intro': 'Combina desviación estándar y conflicto entre criterios.',
        'equations': [
            ('Desviación', r'\sigma_{j}=\mathrm{std}(X_{j})'),
            ('Conflicto', r'C_{j}=\sum_k (1-\mathrm{corr}(j,k))'),
            ('Peso', r'w_{j}=\dfrac{\sigma_{j}\,C_{j}}{\sum_k (\sigma_{k}\,C_{k})}'),
        ],
    },
}

MADM_DOCS: dict[str, dict[str, Any]] = {
    'topsis': {
        'title': 'TOPSIS',
        'intro': (
            'Ordena las alternativas según su cercanía relativa a la solución '
            'ideal positiva y su lejanía de la ideal negativa.'
        ),
        'equations': [
            ('Distancia ideal+', r'D_{i}^{+}=\sqrt{\sum_j (v_{ij}-v_{j}^{+})^{2}}'),
            ('Distancia ideal−', r'D_{i}^{-}=\sqrt{\sum_j (v_{ij}-v_{j}^{-})^{2}}'),
            ('Cercanía relativa', r'C_{i}=\dfrac{D_{i}^{-}}{D_{i}^{+}+D_{i}^{-}}'),
        ],
    },
    'wsm': {
        'title': 'WSM (Weighted Sum Model)',
        'intro': 'Suma ponderada de las utilidades normalizadas por dimensión.',
        'equations': [
            ('Puntuación', r'S_{i}=\sum_j w_{j}\,r_{ij}'),
        ],
    },
    'moora': {
        'title': 'MOORA',
        'intro': 'Ratio multiobjetivo: suma de beneficios menos suma de costos.',
        'equations': [
            (
                'Puntuación',
                r'Y_{i}=\sum_{j\in\mathrm{beneficio}} w_{j} r_{ij}'
                r'-\sum_{j\in\mathrm{costo}} w_{j} r_{ij}',
            ),
        ],
    },
    'vikor': {
        'title': 'VIKOR',
        'intro': 'Compromiso entre utilidad de grupo y arrepentimiento individual.',
        'equations': [
            (
                'Índice S',
                r'S_{i}=\sum_j w_{j}\dfrac{f_{j}^{*}-f_{ij}}{f_{j}^{*}-f_{j}^{-}}',
            ),
            (
                'Índice R',
                r'R_{i}=\max_j\left[w_{j}\dfrac{f_{j}^{*}-f_{ij}}{f_{j}^{*}-f_{j}^{-}}\right]',
            ),
        ],
    },
}


def _progress(cb: Callable[[int, str], None] | None, percent: int, stage: str) -> None:
    if cb:
        cb(percent, stage)


def _add_paragraph(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, color=RGBColor(0x33, 0x41, 0x55))
    run.italic = italic


def _add_method_block(doc: Document, doc_info: dict[str, Any] | None, fallback_name: str) -> None:
    if not doc_info:
        _add_paragraph(doc, f'Método aplicado: {fallback_name}.')
        return
    _add_paragraph(doc, doc_info.get('intro') or f'Método: {doc_info.get("title", fallback_name)}.')
    for label, formula in doc_info.get('equations') or []:
        add_eq_line(doc, label, formula, latex=True)


def _matrix_to_rows(
    matrix: list[list[Any]] | None,
    row_labels: list[str] | None,
    col_labels: list[str] | None,
) -> tuple[list[str], list[list[str]]]:
    cols = list(col_labels or [])
    labels = list(row_labels or [])
    headers = (
        ['Alternativa', *cols]
        if cols
        else [
            'Alternativa',
            *(f'C{i + 1}' for i in range(len((matrix or [[]])[0]) if matrix else 0)),
        ]
    )
    rows: list[list[str]] = []
    for i, row in enumerate(matrix or []):
        name = labels[i] if i < len(labels) else f'Alt {i + 1}'
        rows.append([name, *[_fmt_num(v) for v in row]])
    return headers, rows


def _dimensiones_from_resultado(resultado: dict[str, Any]) -> list[dict[str, Any]]:
    alts = resultado.get('alternativas') or []
    template = (alts[0].get('dimensiones') if alts else None) or []
    dims = []
    for d in template:
        dims.append({
            'id': d.get('omoe_id'),
            'nombre': d.get('omoe_nombre') or f"Dimensión {d.get('omoe_id')}",
            'rama': (d.get('rama_evaluacion') or 'omoe').lower(),
        })
    return dims


def _alt_display(alt: dict[str, Any]) -> str:
    nombre = (alt.get('nombre') or '').strip()
    apodo = (alt.get('apodo') or '').strip()
    if apodo and nombre:
        return f'{nombre} ({apodo})'
    return nombre or apodo or f"Alt {alt.get('id', '—')}"


def _active_dims(
    dims: list[dict[str, Any]],
    opciones: dict[str, Any],
) -> list[dict[str, Any]]:
    """Dimensiones realmente usadas por el pipeline (`dimensiones_normalizar`).

    El cálculo (Pareto, normalización y MADM) opera solo sobre las dimensiones
    seleccionadas por el usuario. Si el informe usara todas las dimensiones de la
    plantilla, el filtro de Pareto podría excluir alternativas distintas a las de
    la interfaz. Restringimos al mismo subconjunto y respetamos su presencia.
    """
    dims_norm = opciones.get('dimensiones_normalizar')
    if isinstance(dims_norm, (list, tuple)) and dims_norm:
        wanted = [str(n) for n in dims_norm]
        by_name = {d['nombre']: d for d in dims}
        active = [by_name[n] for n in wanted if n in by_name]
        if active:
            return active
    return list(dims)


def _pareto_inputs(
    resultado: dict[str, Any],
    dims: list[dict[str, Any]],
    opciones: dict[str, Any],
) -> tuple[list[str], list[list[float]], list[str], list[str]]:
    """Reconstruye (alt_names, matriz, dim_names, directions) idéntico al pipeline.

    - Solo dimensiones activas (`dimensiones_normalizar`).
    - Direcciones resueltas por nombre de dimensión (no por índice), igual que el
      pipeline, para que la dominancia coincida con la de la interfaz.
    - Matriz construida por omoe_id (independiente del orden de columnas guardado).
    """
    active_dims = _active_dims(dims, opciones)
    dim_names = [d['nombre'] for d in active_dims]
    directions = resolve_directions_for_dims(active_dims, opciones)

    # Orden estable por identificador de alternativa (orden original), para que la
    # sección de Pareto coincida con el de la interfaz y no con el ranking MADM.
    alts = list(resultado.get('alternativas') or [])

    def _alt_sort_key(alt: dict[str, Any]) -> tuple[int, float, str]:
        raw_id = alt.get('id')
        try:
            return (0, float(raw_id), '')
        except (TypeError, ValueError):
            return (1, 0.0, str(raw_id))

    alts.sort(key=_alt_sort_key)
    alt_names: list[str] = []
    matrix: list[list[float]] = []
    for alt in alts:
        by_id = {
            d.get('omoe_id'): d.get('valor')
            for d in (alt.get('dimensiones') or [])
        }
        row: list[float] = []
        ok = True
        for d in active_dims:
            v = by_id.get(d['id'])
            if v is None:
                ok = False
                break
            row.append(float(v))
        if not ok:
            continue
        alt_names.append(_alt_display(alt))
        matrix.append(row)
    return alt_names, matrix, dim_names, directions


def _build_entrada_matrix(resultado: dict[str, Any]) -> tuple[list[str], list[list[str]]]:
    matriz = resultado.get('matriz_original')
    dims = _dimensiones_from_resultado(resultado)
    dim_names = [d['nombre'] for d in dims]
    alts = resultado.get('alternativas') or []
    if matriz:
        labels = [_alt_display(a) for a in alts] if alts else None
        return _matrix_to_rows(matriz, labels, dim_names or None)

    headers = ['Alternativa', *[d['nombre'] for d in dims]]
    rows = []
    for alt in alts:
        by_id = {
            d.get('omoe_id'): d.get('valor')
            for d in (alt.get('dimensiones') or [])
        }
        rows.append([
            _alt_display(alt),
            *[_fmt_num(by_id.get(d['id'])) for d in dims],
        ])
    return headers, rows


def _puntos_grafico(resultado: dict[str, Any]) -> list[dict[str, Any]]:
    alts = resultado.get('alternativas') or []
    aplicar_pareto = bool((resultado.get('opciones_calculo') or {}).get('aplicar_pareto'))
    labels = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    puntos = []
    for idx, alt in enumerate(alts):
        valores = {}
        for d in alt.get('dimensiones') or []:
            if d.get('omoe_id') is not None and d.get('valor') is not None:
                valores[d['omoe_id']] = float(d['valor'])
        puntos.append({
            'id': alt.get('id'),
            'nombre': _alt_display(alt),
            'label': labels[idx] if idx < len(labels) else str(idx + 1),
            'valores': valores,
            'overall': alt.get('score_madm') if alt.get('score_madm') is not None else alt.get('valor_global'),
            'score_madm': alt.get('score_madm'),
            'valor_global': alt.get('valor_global'),
            'ranking': alt.get('ranking'),
            'excluida_pareto': bool(alt.get('excluida_pareto')),
        })
    if aplicar_pareto:
        vivos = [p for p in puntos if not p['excluida_pareto']]
        return vivos or puntos
    return puntos


def build_informe_resultados_docx(
    proyecto: Proyecto,
    *,
    historial: SimulacionHistorial | None = None,
    resultado: dict[str, Any] | None = None,
    nombre_calculo: str = '',
    progress_callback: Callable[[int, str], None] | None = None,
) -> bytes:
    """Genera el Informe de resultados (Etapa 3) para un cálculo concreto."""
    progress = lambda p, s: _progress(progress_callback, p, s)  # noqa: E731

    if historial is not None:
        resultado = historial.resultado or {}
        nombre_calculo = (
            (historial.nombre or historial.titulo or '').strip()
            or nombre_calculo
            or f'Cálculo #{historial.id}'
        )
        fecha = historial.fecha_creacion.isoformat() if historial.fecha_creacion else ''
        historial_id = historial.id
    else:
        resultado = resultado or {}
        fecha = ''
        historial_id = resultado.get('historial_id')

    if not resultado:
        raise ValueError('No hay resultado de cálculo para exportar.')

    opciones = resultado.get('opciones_calculo') or {}
    progress(5, 'Preparando documento')

    doc = Document()
    _setup_document_fonts(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    _add_header_logos(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Informe de resultados')
    _set_run_font(run, bold=True, color=RGBColor(0x0F, 0x2C, 0x59))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(proyecto.nombre)
    _set_run_font(run, bold=True)

    progress(12, 'Identificación del cálculo')
    _add_heading(doc, '3. Identificación del cálculo', level=1)
    _add_paragraph(
        doc,
        'Este informe documenta paso a paso el experimento/cálculo seleccionado '
        'en el módulo de simulaciones (Etapa 3), con método científico, matrices '
        'intermedias y gráficos aplicables.',
    )
    id_rows = [
        ['Proyecto', proyecto.nombre],
        ['Nombre del cálculo', nombre_calculo or '—'],
        ['Identificador', str(historial_id) if historial_id else '—'],
        ['Fecha', fecha or '—'],
        ['Configuración', resumen_opciones_calculo(opciones) or '—'],
        [
            'Ganador',
            _fmt(
                (historial.ganador_nombre if historial else None)
                or resultado.get('madm', {}).get('best_alternative')
            ),
        ],
    ]
    _add_table(
        doc,
        ['Campo', 'Valor'],
        id_rows,
        title='Identificación del cálculo',
    )

    # --- 3.1 Dimensiones ---
    progress(18, 'Configuración del pipeline')
    _add_heading(doc, '3.1. Dimensiones activas y sentidos de preferencia', level=2)
    dims = _dimensiones_from_resultado(resultado)
    directions = resolve_directions_for_dims(dims, opciones)
    dim_rows = []
    for d, dir_val in zip(dims, directions):
        sentido_txt = (
            'Min (menor es mejor)' if dir_val == 'min' else 'Max (mayor es mejor)'
        )
        dim_rows.append([d['nombre'], d['rama'].upper(), sentido_txt, 'Sí'])

    _add_table(
        doc,
        ['Dimensión', 'Rama', 'Sentido de preferencia', 'Activa'],
        dim_rows or [['—', '—', '—', '—']],
        title='Dimensiones del cálculo',
    )

    # --- 3.2 Matriz de entrada ---
    progress(28, 'Matriz de utilidades')
    _add_heading(doc, '3.2. Matriz de utilidades (entrada)', level=2)
    _add_paragraph(
        doc,
        'Matriz bruta del cálculo: utilidad agregada por alternativa y dimensión '
        'antes de Pareto, normalización y ranking MADM.',
    )
    headers, rows = _build_entrada_matrix(resultado)
    _add_table(
        doc,
        headers,
        rows or [['—']],
        title='Matriz de utilidades del cálculo',
    )

    # --- 3.3 Pareto (solo si se aplicó; desarrollo milimétrico) ---
    aplicar_pareto = bool(opciones.get('aplicar_pareto'))
    if aplicar_pareto:
        progress(38, 'Filtro Pareto')
        _add_heading(doc, '3.3. Filtro de Pareto', level=2)
        _add_paragraph(
            doc,
            'A continuación se reconstruye el filtro de no dominancia a partir de '
            'la matriz de utilidades: orientación por sentido de preferencia, '
            'sustitución celda a celda y comparaciones que justifican cada exclusión.',
        )
        full_names, full_matrix, pareto_dim_names, pareto_directions = _pareto_inputs(
            resultado, dims, opciones,
        )
        from .pareto_solver import DEFAULT_PARETO_EPSILON

        eps_raw = opciones.get('pareto_epsilon')
        try:
            eps_val = (
                float(eps_raw)
                if eps_raw is not None and eps_raw != ''
                else DEFAULT_PARETO_EPSILON
            )
        except (TypeError, ValueError):
            eps_val = DEFAULT_PARETO_EPSILON
        if full_matrix:
            add_pareto_trace(
                doc,
                matrix=full_matrix,
                alt_names=full_names,
                dim_names=pareto_dim_names,
                directions=pareto_directions,
                epsilon=eps_val,
            )
        else:
            _add_paragraph(
                doc,
                'No fue posible reconstruir la matriz de entrada para el desarrollo de Pareto.',
                italic=True,
            )

    # --- 3.4 Normalización ---
    progress(50, 'Normalización')
    _add_heading(doc, '3.4. Normalización', level=2)
    norm = resultado.get('normalizacion') or {}
    metodo_norm = (
        norm.get('method')
        or opciones.get('normalizacion_metodo')
        or '—'
    )
    doc_norm = NORMALIZATION_DOCS.get(str(metodo_norm))
    _add_heading(
        doc,
        f'3.4.1. Método: {(doc_norm or {}).get("title", metodo_norm)}',
        level=3,
        keep_with_next=True,
    )
    _add_method_block(doc, doc_norm, str(metodo_norm))

    work_names, work_matrix, _ = extract_decision_matrix(
        resultado, dims, only_pareto=True,
    )
    dim_names_work = [d['nombre'] for d in dims]
    if work_matrix:
        add_normalization_trace(
            doc,
            method=str(metodo_norm),
            matrix=work_matrix,
            alt_names=work_names,
            dim_names=dim_names_work,
            directions=directions,
        )

    norm_matrix = norm.get('normalized_matrix')
    norm_alts = norm.get('pareto_alternatives') or work_names or [
        _alt_display(a)
        for a in (resultado.get('alternativas') or [])
        if not a.get('excluida_pareto')
    ]
    norm_dims = norm.get('dimensions') or dim_names_work
    if norm_matrix:
        h, r = _matrix_to_rows(norm_matrix, list(norm_alts), list(norm_dims))
        _add_table(doc, h, r, title='Matriz normalizada')
    else:
        _add_paragraph(doc, 'No hay matriz normalizada almacenada en este cálculo.')

    # --- 3.5 Pesos ---
    progress(62, 'Pesos por dimensión')
    _add_heading(doc, '3.5. Pesos por dimensión', level=2)
    pesos = resultado.get('pesos') or {}
    metodo_pesos = pesos.get('method') or opciones.get('metodo_pesos') or '—'
    doc_pesos = WEIGHT_DOCS.get(str(metodo_pesos))
    _add_heading(
        doc,
        f'3.5.1. Método: {(doc_pesos or {}).get("title", _PESOS_LABELS.get(str(metodo_pesos), metodo_pesos))}',
        level=3,
        keep_with_next=True,
    )
    _add_method_block(doc, doc_pesos, str(metodo_pesos))

    weights_by_dim = pesos.get('weights_by_dimension') or {}
    weights_list = pesos.get('weights') or []
    if isinstance(weights_list, list):
        weights_float = [float(w) for w in weights_list]
    else:
        weights_float = []
    if not weights_float and isinstance(weights_by_dim, dict) and weights_by_dim:
        for d in dims:
            if d['nombre'] in weights_by_dim:
                weights_float.append(float(weights_by_dim[d['nombre']]))
            else:
                weights_float.append(0.0)

    pesos_usuario = opciones.get('pesos_usuario') or pesos.get('user_percentages')
    add_weights_trace(
        doc,
        method=str(metodo_pesos),
        weights=weights_float,
        dim_names=list(norm_dims) if norm_dims else dim_names_work,
        norm_matrix=[
            [float(v) for v in row] for row in (norm_matrix or [])
        ] or None,
        pesos_usuario=(
            [float(p) for p in pesos_usuario]
            if isinstance(pesos_usuario, (list, tuple))
            else None
        ),
    )

    peso_rows = []
    if isinstance(weights_by_dim, dict) and weights_by_dim:
        for name, w in weights_by_dim.items():
            pct = float(w) * 100 if w is not None else None
            peso_rows.append([
                name,
                _fmt_num(w, 6),
                _fmt_num(pct, 2) + ' %' if pct is not None else '—',
            ])
    elif weights_float and dims:
        for i, w in enumerate(weights_float):
            name = dims[i]['nombre'] if i < len(dims) else f'Dimensión {i + 1}'
            pct = float(w) * 100 if w is not None else None
            peso_rows.append([
                name,
                _fmt_num(w, 6),
                _fmt_num(pct, 2) + ' %' if pct is not None else '—',
            ])

    if peso_rows:
        _add_table(
            doc,
            ['Dimensión', 'Peso (w)', 'Peso (%)'],
            peso_rows,
            title='Pesos resultantes por dimensión',
        )
    else:
        _add_paragraph(doc, 'No hay pesos almacenados en este cálculo.')

    # --- 3.6 MADM ---
    progress(74, 'Ranking MADM')
    _add_heading(doc, '3.6. Método MADM y ranking', level=2)
    madm = resultado.get('madm') or {}
    metodo_madm = opciones.get('metodo_madm') or madm.get('method') or '—'
    label_madm = _MADM_LABELS.get(str(metodo_madm), str(metodo_madm))
    doc_madm = MADM_DOCS.get(str(metodo_madm))
    _add_heading(
        doc,
        f'3.6.1. Método: {label_madm}',
        level=3,
        keep_with_next=True,
    )
    _add_method_block(doc, doc_madm, label_madm)

    madm_norm = [
        [float(v) for v in row] for row in (norm_matrix or [])
    ]
    madm_alts = list(norm_alts) if norm_alts else work_names
    madm_dims = list(norm_dims) if norm_dims else dim_names_work
    if madm_norm and weights_float:
        add_madm_trace(
            doc,
            method=str(metodo_madm),
            norm_matrix=madm_norm,
            weights=weights_float,
            alt_names=madm_alts,
            dim_names=madm_dims,
        )

    ranking_rows_table = []
    ranking_for_chart: list[tuple[str, float]] = []
    alts = resultado.get('alternativas') or []
    ranked = sorted(
        [a for a in alts if not a.get('excluida_pareto')],
        key=lambda a: (
            a.get('ranking') is None,
            a.get('ranking') if a.get('ranking') is not None else 9999,
            -(a.get('score_madm') or a.get('valor_global') or 0),
        ),
    )
    if not ranked:
        ranked = list(alts)
    for alt in ranked:
        score = alt.get('score_madm')
        if score is None:
            score = alt.get('valor_global')
        ranking_rows_table.append([
            _fmt(alt.get('ranking')),
            _alt_display(alt),
            _fmt_num(score, 6),
            'Excluida Pareto' if alt.get('excluida_pareto') else 'Activa',
        ])
        if score is not None and not alt.get('excluida_pareto'):
            ranking_for_chart.append((_alt_display(alt), float(score)))

    scores_by_alt = madm.get('scores_by_alternative') or madm.get('preference_by_alternative') or {}
    if isinstance(scores_by_alt, dict) and scores_by_alt and not ranking_for_chart:
        ranking_for_chart = [(k, float(v)) for k, v in scores_by_alt.items()]
        ranking_for_chart.sort(key=lambda x: -x[1])
        ranking_rows_table = [
            [str(i + 1), name, _fmt_num(score, 6), 'Activa']
            for i, (name, score) in enumerate(ranking_for_chart)
        ]

    _add_table(
        doc,
        ['Ranking', 'Alternativa', 'Puntuación', 'Estado'],
        ranking_rows_table or [['—', '—', '—', '—']],
        title=f'Ranking final — {label_madm}',
    )

    bar = ranking_bar_png(ranking_for_chart)
    if bar:
        add_picture_block(
            doc,
            bar,
            f'Figura. Ranking {label_madm} del cálculo «{nombre_calculo}».',
        )

    # --- 3.7 Gráficos ---
    progress(86, 'Gráficos de resultados')
    _add_heading(doc, '3.7. Gráficos de resultados', level=2)
    _add_paragraph(
        doc,
        'Se incluyen gráficos unidimensionales, bidimensionales, tridimensionales '
        'y de radar según las dimensiones activas del cálculo.',
    )
    puntos = _puntos_grafico(resultado)
    dim_metas = _dimensiones_from_resultado(resultado)

    # 3.7.1 1D
    _add_heading(doc, '3.7.1. Gráficos unidimensionales', level=3)
    if not dim_metas:
        _add_paragraph(doc, 'No hay dimensiones para gráficos 1D.', italic=True)
    else:
        for d in dim_metas:
            for horizontal, orient_label in ((True, 'horizontal'), (False, 'vertical')):
                png = bar_1d_png(
                    puntos,
                    d['id'],
                    d['nombre'],
                    horizontal=horizontal,
                )
                if png:
                    add_picture_block(
                        doc,
                        png,
                        f'Figura. Barras {orient_label} — {d["nombre"]}.',
                    )
                else:
                    _add_paragraph(
                        doc,
                        f'No se pudo generar el gráfico 1D {orient_label} de «{d["nombre"]}».',
                        italic=True,
                    )

    # 3.7.2 2D
    _add_heading(doc, '3.7.2. Gráficos bidimensionales (trade-off)', level=3)
    pairs = list(itertools.combinations(dim_metas, 2))
    if not pairs:
        _add_paragraph(doc, 'No hay al menos dos dimensiones para gráficos trade-off.')
    else:
        for a, b in pairs:
            png = scatter_2d_png(
                puntos,
                a['id'],
                b['id'],
                a['nombre'],
                b['nombre'],
                f'{a["nombre"]} vs {b["nombre"]}',
            )
            if png:
                add_picture_block(
                    doc,
                    png,
                    f'Figura. Trade-off {a["nombre"]} × {b["nombre"]}.',
                )
            else:
                _add_paragraph(
                    doc,
                    f'No se pudo generar el gráfico {a["nombre"]} × {b["nombre"]} '
                    '(faltan puntos con ambas coordenadas).',
                    italic=True,
                )

    # 3.7.2b Espacio de decisión 2D con mapa de calor (color = 3ª dimensión)
    _add_heading(
        doc,
        '3.7.2.1. Espacio de decisión 2D (mapa de calor por dimensión)',
        level=3,
    )
    _add_paragraph(
        doc,
        'Dos dimensiones se ubican en los ejes X e Y y una tercera dimensión se '
        'representa como color (barra lateral). Se rotan las dimensiones para que '
        'cada una aparezca como color al menos una vez.',
    )
    if len(dim_metas) < 3:
        _add_paragraph(
            doc,
            'Se requieren al menos tres dimensiones activas para el mapa de calor 2D '
            '(dos ejes + una dimensión de color).',
            italic=True,
        )
    else:
        max_heat = 12
        heat_specs: list[tuple[dict, dict, dict]] = []
        for a, b, c in itertools.combinations(dim_metas, 3):
            # Rotar cuál dimensión va como color (barra) manteniendo las otras dos en los ejes.
            heat_specs.append((a, b, c))  # color = c
            heat_specs.append((a, c, b))  # color = b
            heat_specs.append((b, c, a))  # color = a
        if len(heat_specs) > max_heat:
            _add_paragraph(
                doc,
                f'Hay {len(heat_specs)} combinaciones posibles; se muestran las '
                f'primeras {max_heat}.',
                italic=True,
            )
            heat_specs = heat_specs[:max_heat]

        heat_ok = False
        for x_dim, y_dim, color_dim in heat_specs:
            heat = scatter_2d_heat_png(
                puntos,
                x_dim['id'],
                y_dim['id'],
                color_dim['id'],
                x_dim['nombre'],
                y_dim['nombre'],
                color_dim['nombre'],
                f'{x_dim["nombre"]} × {y_dim["nombre"]} · color: {color_dim["nombre"]}',
            )
            if heat:
                heat_ok = True
                add_picture_block(
                    doc,
                    heat,
                    f'Figura. Espacio de decisión {x_dim["nombre"]} × {y_dim["nombre"]} '
                    f'(color = {color_dim["nombre"]}).',
                )
        if not heat_ok:
            _add_paragraph(
                doc,
                'No se pudo generar el mapa de calor 2D (faltan valores por dimensión).',
                italic=True,
            )

    # 3.7.3 3D
    _add_heading(doc, '3.7.3. Gráficos tridimensionales', level=3)
    triples = list(itertools.combinations(dim_metas, 3))
    if not triples:
        _add_paragraph(
            doc,
            'No hay al menos tres dimensiones para gráficos tridimensionales.',
            italic=True,
        )
    else:
        max_triples = 8
        shown = triples[:max_triples]
        if len(triples) > max_triples:
            _add_paragraph(
                doc,
                f'Hay {len(triples)} combinaciones tridimensionales; '
                f'se muestran las primeras {max_triples}.',
                italic=True,
            )
        for a, b, c in shown:
            png = scatter_3d_png(
                puntos,
                a['id'],
                b['id'],
                c['id'],
                a['nombre'],
                b['nombre'],
                c['nombre'],
                f'{a["nombre"]} × {b["nombre"]} × {c["nombre"]}',
            )
            if png:
                add_picture_block(
                    doc,
                    png,
                    f'Figura. 3D {a["nombre"]} × {b["nombre"]} × {c["nombre"]}.',
                )
            else:
                _add_paragraph(
                    doc,
                    f'No se pudo generar el gráfico 3D '
                    f'{a["nombre"]} × {b["nombre"]} × {c["nombre"]}.',
                    italic=True,
                )

    # 3.7.4 Radar
    _add_heading(doc, '3.7.4. Gráfico de radar', level=3)
    radar = radar_png(puntos, dim_metas)
    if radar:
        add_picture_block(
            doc,
            radar,
            'Figura. Radar de dimensiones del cálculo.',
        )
    else:
        _add_paragraph(
            doc,
            'No se pudo generar el gráfico de radar (se requieren ≥2 dimensiones).',
            italic=True,
        )

    # 3.7.5 Línea por ranking: Utilidad vs puntuación MADM
    _add_heading(
        doc,
        f'3.7.5. Línea por ranking: Utilidad vs {label_madm}',
        level=3,
    )
    _add_paragraph(
        doc,
        'Compara, por alternativa (ordenadas por ranking), la utilidad agregada del '
        f'árbol de criterios frente a la puntuación del método {label_madm}.',
    )
    linea = ranking_util_madm_png(puntos, madm_label=label_madm)
    if linea:
        add_picture_block(
            doc,
            linea,
            f'Figura. Ranking — Utilidad (valor global) vs puntuación {label_madm}.',
        )
    else:
        _add_paragraph(
            doc,
            'No se pudo generar la línea Utilidad vs MADM '
            '(se requieren utilidad y puntuación por alternativa).',
            italic=True,
        )

    progress(98, 'Guardando documento')
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Documento generado por HATD · Informe de resultados (por cálculo)')
    _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

    buf = BytesIO()
    doc.save(buf)
    progress(99, 'Preparando descarga')
    return buf.getvalue()


def build_informe_resultados_from_historial(
    proyecto: Proyecto,
    historial_id: int,
    *,
    progress_callback: Callable[[int, str], None] | None = None,
) -> bytes:
    row = SimulacionHistorial.objects.get(pk=historial_id, proyecto=proyecto)
    return build_informe_resultados_docx(
        proyecto,
        historial=row,
        progress_callback=progress_callback,
    )
