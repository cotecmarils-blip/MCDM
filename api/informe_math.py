"""Ecuaciones nativas de Word (OMML) a partir de LaTeX.

Flujo: LaTeX → MathML (latex2mathml) → OMML (mathml2omml) → párrafo python-docx.
Quedan como ecuaciones editables de Word, no como imágenes.
"""
from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from math2docx import add_math

_LABEL_COLOR = RGBColor(0x0F, 0x2C, 0x59)
_BODY_COLOR = RGBColor(0x1E, 0x29, 0x3B)

# Sustituciones Unicode → LaTeX (ecuaciones legacy del informe de proyecto).
_UNICODE_TO_LATEX = (
    ('Σ', r'\sum '),
    ('∑', r'\sum '),
    ('√', r'\sqrt'),
    ('·', r'\cdot '),
    ('×', r'\times '),
    ('−', '-'),
    ('–', '-'),
    ('—', '-'),
    ('…', r'\ldots '),
    ('≤', r'\leq '),
    ('≥', r'\geq '),
    ('≠', r'\neq '),
    ('≈', r'\approx '),
    ('∞', r'\infty '),
    ('ε', r'\varepsilon '),
    ('ϵ', r'\epsilon '),
    ('σ', r'\sigma '),
    ('μ', r'\mu '),
    ('π', r'\pi '),
    ('α', r'\alpha '),
    ('β', r'\beta '),
    ('γ', r'\gamma '),
    ('δ', r'\delta '),
    ('Δ', r'\Delta '),
    ('⇒', r'\Rightarrow '),
    ('→', r'\rightarrow '),
    ('∈', r'\in '),
    ('∀', r'\forall '),
    ('∃', r'\exists '),
    ('±', r'\pm '),
    ('°', r'^{\circ}'),
    ('²', r'^{2}'),
    ('³', r'^{3}'),
    ('₁', r'_{1}'),
    ('₂', r'_{2}'),
    ('ᵢ', r'_{i}'),
    ('ⱼ', r'_{j}'),
    ('ₖ', r'_{k}'),
    ('ₘ', r'_{m}'),
    ('ₙ', r'_{n}'),
    ('ₐ', r'_{a}'),
    ('ᵈ', r'_{d}'),
)


def _set_run_font(run, *, bold: bool = False, color: RGBColor | None = None, size_pt: float = 11) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = 'Arial'
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn('w:ascii'), 'Arial')
    r_fonts.set(qn('w:hAnsi'), 'Arial')
    r_fonts.set(qn('w:cs'), 'Arial')


def tex_text(value: object) -> str:
    """Texto literal seguro dentro de \\text{...} para subíndices con nombres."""
    text = str(value or '')
    text = (
        text.replace('\\', '/')
        .replace('{', '(')
        .replace('}', ')')
        .replace('_', '-')
        .replace('^', '')
        .replace('%', ' pct')
        .replace('&', ' y ')
        .replace('#', '')
        .replace('$', '')
    )
    return rf'\text{{{text}}}'


def normalize_to_latex(formula: str) -> str:
    """Normaliza pseudo-LaTeX / Unicode a LaTeX usable por latex2mathml."""
    expr = (formula or '').strip()
    if not expr:
        return expr
    # Quitar delimitadores $...$ si vienen envueltos.
    if expr.startswith('$$') and expr.endswith('$$'):
        expr = expr[2:-2].strip()
    elif expr.startswith('$') and expr.endswith('$'):
        expr = expr[1:-1].strip()

    for src, dst in _UNICODE_TO_LATEX:
        expr = expr.replace(src, dst)

    # "ó" / "o" entre alternativas → \mathrm{o}
    expr = re.sub(r'\s+ó\s+', r'\\ \\mathrm{o}\\ ', expr)
    # Espacios múltiples
    expr = re.sub(r'[ \t]+', ' ', expr).strip()
    return expr


def latex_to_omml_element(latex: str):
    """Convierte LaTeX a un elemento OMML listo para insertar en un párrafo."""
    expr = normalize_to_latex(latex)
    mathml_output = __import__('latex2mathml.converter', fromlist=['convert']).convert(expr)
    omml_output = __import__('mathml2omml', fromlist=['convert']).convert(mathml_output)
    xml_output = (
        '<p xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'{omml_output}</p>'
    )
    return parse_xml(xml_output)[0]


def _append_omml(paragraph, latex: str) -> bool:
    """Inserta OMML en el párrafo. True si tuvo éxito."""
    try:
        add_math(paragraph, normalize_to_latex(latex))
        return True
    except Exception:
        try:
            paragraph._p.append(latex_to_omml_element(latex))
            return True
        except Exception:
            return False


def _fallback_plain(paragraph, formula: str) -> None:
    run = paragraph.add_run(formula)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    try:
        run.font.color.rgb = _BODY_COLOR
    except Exception:
        pass


def add_latex_equation(
    doc: Document,
    latex: str,
    *,
    label: str | None = None,
    number: int | None = None,
    center: bool = True,
) -> None:
    """Inserta una ecuación OMML nativa (opcionalmente con etiqueta y número)."""
    if label:
        p_label = doc.add_paragraph()
        run = p_label.add_run(f'{label}:')
        _set_run_font(run, bold=True, color=_LABEL_COLOR)

    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True

    ok = _append_omml(p, latex)
    if not ok:
        _fallback_plain(p, latex)

    if number is not None:
        run_n = p.add_run(f'    ({number})')
        _set_run_font(run_n, color=_BODY_COLOR, size_pt=11)


def add_eq_line(
    doc: Document,
    label: str,
    formula: str,
    *,
    latex: bool = True,
) -> None:
    """Línea «etiqueta + fórmula». Si latex=False, texto plano (sustituciones)."""
    if latex:
        add_latex_equation(doc, formula, label=label)
        return
    p = doc.add_paragraph()
    run_l = p.add_run(f'{label}: ')
    _set_run_font(run_l, bold=True, color=_LABEL_COLOR)
    run_f = p.add_run(formula)
    _set_run_font(run_f, color=_BODY_COLOR)
