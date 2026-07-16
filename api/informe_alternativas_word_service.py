"""Informe Word de las alternativas de un proyecto (con fotos y anexos).

Documento independiente del módulo de Alternativas: ficha por alternativa con
nombre, apodo, descripción, referencia, costo, capacidades, características,
foto e inventario de documentos anexos.
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, RGBColor

from .informe_word_service import (
    _add_header_logos,
    _add_heading,
    _add_table,
    _set_run_font,
    _setup_document_fonts,
)
from .models import Alternativa, Proyecto

_BODY_COLOR = RGBColor(0x1E, 0x29, 0x3B)
_MUTED_COLOR = RGBColor(0x47, 0x55, 0x69)
_TITLE_COLOR = RGBColor(0x0F, 0x2C, 0x59)
_ERROR_COLOR = RGBColor(0xB9, 0x1C, 0x1C)

# Caracteres de control no válidos en XML 1.0 (excepto tab, salto de línea y
# retorno de carro). Texto ingresado por el usuario puede traerlos al pegar
# desde PDFs u otras fuentes, y python-docx lanza ValueError al escribirlos.
_ILLEGAL_XML_CHARS = re.compile(
    '[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ud800-\udfff\ufffe\uffff]'
)


def _xs(value: Any, empty: str = '—') -> str:
    """Convierte a str y elimina caracteres no válidos para XML/Word."""
    if value is None:
        return empty
    text = _ILLEGAL_XML_CHARS.sub('', str(value))
    return text if text.strip() else empty


def _p(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
       color: RGBColor | None = None, size_pt: float = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_xs(text, empty=''))
    _set_run_font(run, bold=bold, color=color, size_pt=size_pt)
    run.italic = italic


def _costo_txt(alt: Alternativa) -> str:
    if alt.costo is None:
        return '—'
    return _xs(f'{alt.costo} {alt.costo_unidad}')


def _add_foto(doc: Document, alt: Alternativa) -> None:
    if not alt.foto or not alt.foto.name:
        return
    try:
        foto_path = Path(alt.foto.path)
    except Exception:
        return
    if not foto_path.is_file():
        return
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.keep_with_next = True
    try:
        pic.add_run().add_picture(str(foto_path), height=Cm(6.5))
    except Exception:
        _p(doc, 'No se pudo incorporar la imagen de la alternativa.', color=_ERROR_COLOR)


def build_informe_alternativas_docx(proyecto: Proyecto) -> bytes:
    """Genera el informe Word de las alternativas del proyecto (con fotos)."""
    alternativas = list(
        Alternativa.objects.filter(proyecto=proyecto)
        .prefetch_related('capacidades', 'caracteristicas__plantilla', 'documentos')
        .order_by('id')
    )

    doc = Document()
    _setup_document_fonts(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    _add_header_logos(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Informe de alternativas')
    _set_run_font(run, bold=True, color=_TITLE_COLOR)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(_xs(f'Proyecto: {proyecto.nombre}'))
    _set_run_font(run, bold=True)

    _p(
        doc,
        f'Total de alternativas: {len(alternativas)}.',
        color=_MUTED_COLOR,
    )

    if not alternativas:
        _p(doc, 'Este proyecto no tiene alternativas registradas.', italic=True)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Tabla resumen general.
    resumen_rows = []
    for alt in alternativas:
        nombre = _xs(alt.nombre)
        if alt.apodo:
            nombre = f'{nombre} ({_xs(alt.apodo)})'
        resumen_rows.append([nombre, _costo_txt(alt), str(alt.caracteristicas.count()),
                             str(alt.capacidades.count()), str(alt.documentos.count())])
    _add_heading(doc, '1. Resumen de alternativas', level=1, keep_with_next=True)
    _add_table(
        doc,
        ['Alternativa', 'Costo declarado', 'N.º características', 'N.º capacidades', 'N.º documentos'],
        resumen_rows,
        title='Resumen general de alternativas',
        new_page=False,
    )

    # Ficha detallada por alternativa.
    _add_heading(doc, '2. Ficha por alternativa', level=1)
    for idx, alt in enumerate(alternativas, start=1):
        nombre = _xs(alt.nombre)
        if alt.apodo:
            nombre = f'{nombre} ({_xs(alt.apodo)})'
        _add_heading(doc, f'2.{idx}. {nombre}', level=2, keep_with_next=True)

        _add_foto(doc, alt)

        info_rows = [
            ['Nombre', _xs(alt.nombre)],
            ['Apodo', _xs(alt.apodo)],
            ['Costo declarado', _costo_txt(alt)],
            ['Descripción', _xs(alt.descripcion)],
            ['Referencia', _xs(alt.referencia)],
        ]
        _add_table(
            doc,
            ['Campo', 'Valor'],
            info_rows,
            title=f'Datos generales — {nombre}',
            new_page=False,
        )

        capacidades = list(alt.capacidades.all())
        if capacidades:
            cap_rows = [[_xs(c.nombre), _xs(c.descripcion)] for c in capacidades]
            _add_table(
                doc,
                ['Capacidad', 'Descripción'],
                cap_rows,
                title=f'Capacidades — {nombre}',
                new_page=False,
            )

        caracteristicas = sorted(
            alt.caracteristicas.all(),
            key=lambda c: (getattr(c.plantilla, 'orden', 0) or 0, c.id),
        )
        if caracteristicas:
            car_rows = [
                [
                    _xs(c.plantilla.nombre),
                    _xs(c.dato),
                    _xs(c.plantilla.unidad),
                ]
                for c in caracteristicas
            ]
            _add_table(
                doc,
                ['Característica', 'Valor', 'Unidad'],
                car_rows,
                title=f'Características — {nombre}',
                new_page=False,
            )

        documentos = list(alt.documentos.all())
        if documentos:
            doc_rows = [
                [_xs(d.nombre or (Path(d.archivo.name).name if d.archivo else '')),
                 _xs(d.archivo.name if d.archivo else '')]
                for d in documentos
            ]
            _add_table(
                doc,
                ['Documento', 'Archivo'],
                doc_rows,
                title=f'Documentos anexos — {nombre}',
                new_page=False,
            )
        if alt.anexo and alt.anexo.name:
            _p(doc, _xs(f'Anexo principal: {alt.anexo.name}'), color=_MUTED_COLOR, size_pt=9)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Documento generado por HATD · Informe de alternativas')
    _set_run_font(run, color=RGBColor(0x64, 0x74, 0x8B))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
